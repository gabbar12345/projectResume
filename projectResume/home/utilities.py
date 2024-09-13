import tempfile
from docx import Document
from docx.shared import Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
import openai
import ast
from home.constants import pdf_path,image_path
import os
from home.prompts import sy
from fpdf import FPDF
from home.prompts import *

# from home.gcp_secrets import access_secret_version
from projectResume import settings


# openai.api_key =settings.API_KEY

def get_response(user_prompt,system_prompt=sy):
    response = openai.chat.completions.create(
    # model="gpt-4", 
    model="gpt-4o",                    
    messages=[{"role": "system", "content": system_prompt,},
              {"role": "user", "content": user_prompt}],
    temperature=1,
    top_p=1
    )
    # print(response)
    result=response.choices[0].message.content
    return result

def safe_literal_eval(expression):
    try:
        # Attempt to evaluate the string as a Python literal
        result = ast.literal_eval(expression)
        return result
    except (ValueError, SyntaxError):
        # If evaluation fails, return the original string
        return expression

def formatedResponse(prompt):
    Body=get_response(user_prompt=prompt)
    changeBody = safe_literal_eval(Body)
    return changeBody

class PersonalDetails:
    def __init__(self, name, email, phone, linkedin, photo_path):
        self.name = name
        self.email = email
        self.phone = phone
        self.linkedin = linkedin
        self.photo_path = photo_path

class Chapter:
    def __init__(self, title, subtitle, body):
        self.title = title
        self.subtitle = subtitle
        self.body = body

class AcademicDetail:
    def __init__(self, year, degree, institute_university, cgpa_percentage):
        self.year = year
        self.degree = degree
        self.institute_university = institute_university
        self.cgpa_percentage = cgpa_percentage

class Project:
    def __init__(self, title, description):
        self.title = title
        self.description = description

class ResearchPaper:
    def __init__(self, title, authors, publication):
        self.title = title
        self.authors = authors
        self.publication = publication

class ResumePDF(FPDF):
    def __init__(self, personal_details, chapters, academic_details, professional_summary, skills, positions_of_responsibility, projects, research_papers):
        super().__init__()
        self.personal_details = personal_details
        self.academic_details = academic_details
        self.chapters = chapters
        self.professional_summary = professional_summary
        self.skills = skills
        self.positions_of_responsibility = positions_of_responsibility
        self.projects = projects
        self.research_papers = research_papers

    def header(self):
        self.set_fill_color(240, 240, 240)  # Lighter grey background
        self.rect(0, 0, 210, 48, 'F')  # Grey rectangle for header
        self.set_font('Arial', 'B', 24)
        self.set_text_color(0, 0, 100)  # Dark blue color for name
        self.cell(150, 15, self.personal_details.name, 0, 1, 'L')
        self.set_font('Arial', 'I', 12)
        self.set_text_color(100, 100, 100)  # Grey color for subtext
        self.cell(150, 7, "Software Engineer | Full-Stack Developer", 0, 1, 'L')
        self.image(self.personal_details.photo_path, 170, 5, 30, 30)
        self.ln(0)
        self.set_font('Arial', '', 10)
        self.set_text_color(0, 0, 0)  # Reset text color to black
        self.cell(150, 5, self.personal_details.email, 0, 1, 'L')
        self.cell(150, 5, self.personal_details.phone, 0, 1, 'L')
        self.cell(150, 5, self.personal_details.linkedin, 0, 1, 'L', link=self.personal_details.linkedin)
        self.ln(0)
        # Add red horizontal line at the bottom of the grey background
        self.set_draw_color(255, 0, 0)  # Set line color to red
        self.line(0, 49, 210, 49)  # Draw line at y=48 (bottom of grey background)
        self.set_draw_color(0, 0, 0)  # Reset draw color to black

    def footer(self):
        self.set_y(-15)
        self.set_font('Arial', 'I', 8)
        self.cell(0, 10, f'Page {self.page_no()}', 0, 0, 'C')

    def add_section_title(self, title):
        self.set_font('Arial', 'B', 12)
        self.cell(0, 10, title, 'B', 1, 'L')
        self.ln(1)

    def add_chapter(self, chapter):
        self.set_font('Arial', 'B', 11)
        self.cell(130, 6, chapter.title, 0, 0, 'L')
        self.cell(60, 6, chapter.subtitle, 0, 1, 'R')
        self.set_font('Arial', '', 10)
        self.multi_cell(0, 5, chapter.body['description'])
        for point in chapter.body['bullet_points']:
            self.cell(5)
            self.cell(0, 5, chr(149) + ' ' + point, 0, 1)
        self.ln(2)

    def add_academic_details(self):
        self.add_section_title('EDUCATION')
        for detail in self.academic_details:
            self.set_font('Arial', 'B', 10)
            self.cell(130, 6, detail.degree, 0, 0)
            self.cell(60, 6, detail.year, 0, 1, 'R')
            self.set_font('Arial', '', 10)
            self.cell(130, 5, detail.institute_university, 0, 0)
            self.cell(60, 5, detail.cgpa_percentage, 0, 1, 'R')
        self.ln(2)

    def add_professional_summary(self):
        self.add_section_title('PROFESSIONAL SUMMARY')
        self.set_font('Arial', '', 10)
        self.multi_cell(0, 5, self.professional_summary)
        self.ln(1)

    def add_skills(self):
        self.add_section_title('SKILLS')
        self.set_font('Arial', '', 10)
        skills_per_line = 4
        for i in range(0, len(self.skills), skills_per_line):
            self.cell(0, 5, ' | '.join(self.skills[i:i+skills_per_line]), 0, 1)
        self.ln(2)

    def add_position_of_responsibility(self):
        self.add_section_title('POSITIONS OF RESPONSIBILITY')
        for position in self.positions_of_responsibility:
            self.set_font('Arial', 'B', 10)
            self.cell(130, 6, f"{position['position']} - {position['company']}", 0, 0)
            self.cell(60, 6, position['place'], 0, 1, 'R')
            self.set_font('Arial', '', 10)
            self.multi_cell(0, 5, position['impact'])
            self.ln(1)

    def add_projects(self):
        self.add_section_title('PROJECTS')
        for project in self.projects:
            self.set_font('Arial', 'B', 10)
            self.cell(0, 6, project.title, 0, 1)
            self.set_font('Arial', '', 10)
            self.multi_cell(0, 5, project.description)
            self.ln(1)

    def add_research_papers(self):
        self.add_section_title('RESEARCH PAPERS')
        for paper in self.research_papers:
            self.set_font('Arial', 'B', 10)
            self.cell(0, 6, paper.title, 0, 1)
            self.set_font('Arial', '', 10)
            self.cell(0, 5, f"Authors: {paper.authors}", 0, 1)
            self.cell(0, 5, f"Publication: {paper.publication}", 0, 1)
            self.ln(1)

def get_sample_data():
    personal_details = PersonalDetails(
        name="John Doe",
        email="john.doe@example.com",
        phone="+1 (123) 456-7890",
        linkedin="www.linkedin.com/in/johndoe",
        photo_path=image_path
    )

    academic_details = [
        AcademicDetail("2022", "B.S. in Computer Science", "University of Technology", "GPA: 3.8/4.0"),
        AcademicDetail("2018", "High School Diploma", "Central High School", "GPA: 4.0/4.0")
    ]

    chapters = [
        Chapter('Software Engineer', 'XYZ Corp (2020-Present)', {
            'description': 'Lead developer for enterprise-level applications.',
            'bullet_points': [
                'Implemented microservices architecture, improving system efficiency by 40%',
                'Mentored junior developers, increasing team productivity by 25%'
            ]
        })
    ]

    professional_summary = "Dedicated software engineer with 3+ years of experience in developing scalable web applications. Proficient in full-stack development with a focus on cloud technologies. Proven track record of delivering high-quality solutions and driving innovation in fast-paced environments."

    skills = ["Python", "JavaScript", "React", "Node.js", "AWS", "Docker", "Git", "SQL", "MongoDB", "RESTful APIs"]

    positions_of_responsibility = [
        {
            'company': 'Tech Society',
            'place': 'University',
            'position': 'President',
            'impact': 'Organized 10+ tech workshops, increasing member engagement by 50% and facilitating industry connections for students.'
        }
    ]

    projects = [
        Project("AI-powered Chatbot", "Developed a chatbot using natural language processing techniques to improve customer service efficiency."),
        Project("E-commerce Platform", "Built a scalable e-commerce platform using React and Node.js, handling 10,000+ daily active users.")
    ]

    research_papers = [
        ResearchPaper("Machine Learning in Healthcare", "John Doe, Jane Smith", "International Journal of Medical Informatics"),
        ResearchPaper("Blockchain for Supply Chain Management", "John Doe, et al.", "IEEE Transactions on Engineering Management")
    ]

    return personal_details, academic_details, chapters, professional_summary, skills, positions_of_responsibility, projects, research_papers


def create_resume(fields):
    document = Document()

    # Title and Subtitle Formatting
    title_style = document.styles['Heading 1']
    title_font = title_style.font
    title_font.name = 'Arial'
    title_font.size = Pt(24)
    title_font.color.rgb = RGBColor(0, 0, 139)  # Dark Blue

    subtitle_style = document.styles['Heading 2']
    subtitle_font = subtitle_style.font
    subtitle_font.name = 'Arial'
    subtitle_font.size = Pt(14)
    subtitle_font.color.rgb = RGBColor(0, 128, 128)  # Teal

    body_style = document.styles['Normal']
    body_font = body_style.font
    body_font.name = 'Calibri'
    body_font.size = Pt(11)
    body_font.color.rgb = RGBColor(0, 0, 0)  # Black

    # Add Title
    document.add_heading(fields['name'], level=1).alignment = WD_ALIGN_PARAGRAPH.CENTER
    document.add_heading(fields['contact information'], level=2).alignment = WD_ALIGN_PARAGRAPH.CENTER

    # Add Sections
    for section, content in fields['sections'].items():
        document.add_heading(section, level=2)
        for item in content:
            p = document.add_paragraph(item)
    return document



def save_resume(document, file_name=pdf_path):
    # Check if file exists and if so, that we can write to it
    if os.path.exists(file_name):
        if not os.access(file_name, os.W_OK):
            raise PermissionError(f"No write permission for file: {file_name}")

    try:
        # Save to a temporary file first
        temp_file = tempfile.NamedTemporaryFile(delete=False)
        temp_file_name = temp_file.name
        temp_file.close()

        # Saving document to a temporary file
        document.save(temp_file_name)

        # Replace the original file with the temporary file
        os.replace(temp_file_name, file_name)
        print(f"Document successfully saved to {file_name}")

    except PermissionError as e:
        print(f"PermissionError: {e}")
    except Exception as e:
        print(f"An unexpected error occurred: {e}")
        # Clean up temporary file, if it exists
        if os.path.exists(temp_file_name):
            os.remove(temp_file_name)

def get_home_directory():
    return os.path.expanduser("~")

def generate_resume2(jobRole):
    personal_details, academic_details, chapters, professional_summary, skills, positions_of_responsibility, projects, research_papers = get_sample_data()
            # Chapter Prompt Changes
    chapter_prompt=resumePrompt.format(preData='chapters',data=chapters[0].body,job_role=jobRole)
    response=formatedResponse(chapter_prompt)
    chapters[0].body=response

    # Professional Summary Prompt Changes
    Prompt=resumePrompt.format(preData='professionalSummary',data=professional_summary,job_role=jobRole)
    summary_response=formatedResponse(Prompt)
    professional_summary=summary_response

    for i in range(len(projects)):
        prompt=resumePrompt.format(preData='project description',data=projects[i].description,job_role=jobRole)
        project_response=formatedResponse(prompt)
        projects[i].description=project_response

    # Generate PDF
    pdf = ResumePDF(personal_details, chapters, academic_details, professional_summary, skills, positions_of_responsibility, projects, research_papers)
    pdf.add_page()

    # Add light grey background for the entire page
    pdf.set_fill_color(250, 250, 250)  # Even lighter grey background
    pdf.rect(0, 48, 210, 297-40, 'F')  # 297 is A4 height, 40 is header height

    pdf.add_professional_summary()
    for chapter in chapters:
        pdf.add_chapter(chapter)
    pdf.add_academic_details()
    pdf.add_skills()
    pdf.add_position_of_responsibility()
    pdf.add_projects()
    pdf.add_research_papers()

    return pdf
